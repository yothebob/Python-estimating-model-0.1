from docx import Document
from datetime import date
from docx.shared import Pt, Inches,RGBColor
import sys
import rail_description as rd

print('proposal writer loaded...')
def menu():
    enter = input('make proposal? (y/n): ')
    if enter.lower() == 'y':
        write_proposal([40],[60])
    elif enter.lower() == 'n':
        sys.exit()
    else:
        print('please try again')
        menu()



def write_proposal(_bid_lf,_bid_lfprice):
    print('start making proposal...')

    # Document Variables
    customer_name = input('customer rep name?\n :')
    customer_company =input('customer company name?\n :')
    contact_info =input('customer contact info?\n :')
    company_address =input('customer Business address?\n :')
    job_address = input('job site address?\n :')
    job_name = input('Job name?\n :')
    today = date.today()
    d1 = today.strftime("%m/%d/%Y")
    subtotal = 0
    total_info = [d1,customer_name,contact_info,customer_company,company_address,job_name,job_address]
        
    #Document Setup
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(3.4)
    paragraph_format.space_after = Pt(0)
    sections = document.sections
    
    for section in sections:
        section.top_margin = Inches(.5)
        section.bottom_margin = Inches(.4)
        section.left_margin = Inches(.8)
        section.right_margin = Inches(.8)

    #header setup
    header = document.sections[0].header
    h1 = header.paragraphs[0].add_run("\tPrecision Rail of Oregon, LLC\n").font.size = Pt(36)
    h2 = header.paragraphs[0].add_run("10735 SE Foster RD").font.size = Pt(14)
    header.paragraphs[0].add_run('\t').add_picture('Alumarail_logo.png', width = Inches(1.8))
    h3 = header.paragraphs[0].add_run("\tPortland, Oregon 97266").font.size = Pt(14)
    h1.italic = True
    h2.italic = True
    h3.italic = True
    h1.bold = True
    h2.bold = True
    h3.bold = True
    header.paragraphs[0].style.font.color.rgb = RGBColor(0, 96, 43)

    #footer setup
    footer = document.sections[0].footer
    footer.paragraphs[0].add_run('Phone: 503-512-5353\tFax: 503-668-8968\twww.Precisionrail.com').font.size = Pt(14)
    footer.paragraphs[0].style.font.color.rgb = RGBColor(0, 96, 43)

    # start Document
    document.add_paragraph('\n')
    for item in total_info:
        document.add_paragraph(str(item))
    document.add_paragraph('')
    document.add_paragraph('Dear ' + customer_name + ',\n')
    p1 = document.add_paragraph()
    p1.style = document.styles['Normal']
    
    p1.add_run('Precision Rail of Oregon is pleased to provide the following proposal for: ')
    p1.add_run(job_name + ', BUDGET Rev-0 \n\n').bold = True
    p1.add_run('Items furnished by Precision Rail of Oregon: Submittal drawings, engineering, materials, and installation.\n\n').bold = True
    p1.add_run('Submittals:').bold = True
    p1.add_run(' Pricing includes 1 submittal based off plans and 1 revision once corrections are recieved from GC. Any Additional revisions to be billed at 145.00 per hour plus materials and handling. \n\n')

    #sections
    for num in range(len(_bid_lf)):
        bid_area = input('what is this section called? \n: ')
        b1_height,b1_post,b1_mount,b1_top,b1_bottom,b1_infill,b1_space,b1_type = rd.get_description()
        b1 = rd.return_description(b1_height,b1_post,b1_mount,b1_top,b1_bottom,b1_infill,b1_space,b1_type)
        p1.add_run('Bid Item - {} Tall {} ({})\n'.format(b1[0],b1[7],bid_area)).bold = True
        p1.add_run(" {} {}. {}, {} with {}. Posts spacing to be evenly spaced and not exceed {} per engineering and customer request. Support blocking by others. Standard color(Black, Bronze, White). ".format(b1[1],b1[2],b1[3],b1[4],b1[5],b1[6]))    
        if b1[7] == 'Grab rail':
            p1.add_run('Handrails are all ADA Compliant.')
        p1.add_run('\n\n')
        p1.add_run('\t'*7 + ' Total {} LF @ ${}.00 per LF = ${}.00*\n\n\n'.format(str(_bid_lf[num]),str(_bid_lfprice[num]),str(_bid_lf[num]*_bid_lfprice[num]))).bold = True
        subtotal += _bid_lf[num] * _bid_lfprice[num]
        
    p1.add_run('\n'*3)
    p1.add_run('\t'*10 + 'Sub Total = {}.00*\n\n\n'.format(str(subtotal))).bold = True
    
    p1.add_run('\t\t*This price quote is valid for 3 months from the date of this document*\n\n').italic = True

    p1.add_run('Assumptions\n').bold = True
    p1.add_run('The following assumptions were made in support of this estimate:')

    p1.add_run(
                """
        1.	Electrical utilities available on site.
        2.	Sanitation facilities will be provided and available on site.
        3.	Core holes, provided by others, will be cleaned out and ready for post installation.
        4.	Fall restraint anchor points will be available and cleaned out ready for use.
        5.	Paint / PPG Duracron with a 5 year warranty.
                """)
    p1.add_run('\n')
    p1.add_run('Items EXCLUDED by Precision Rail of Oregon unless noted above:')

    p1.add_run(
        """
        1.	Deferred permits or any items not specifically included is considered furnished by others.
        2.	Taxes such as sales, local municipality, gross receipts tax and/or local business licenses.
        3.	Union, prevailing wage and/or workforce training installation
        4.	Insurance requirements above and beyond: $1M/$2M (occurrence/aggregate); and $3M
            umbrella.
        5.	Performance & payment bonds.
        6.	Pollution insurance requirements.
        7.	Deviations from project plans that impede the installation of our rail as planned.
        8.	Marking / locating rebar tensions wires
        9.	Coverage / Protection of any Glazing, Brick, Building materials
        10. Inspection for testing (example UT, NDT & others)
        11. Flaggers, and / or any personnel for traffic control
        12. Lifts, swing stages, cranes, or other equipment required to install are not included in this bid
                and are to be provided by the GC.

        """)
    p1.add_run('\n\n')

    p1.add_run("Submittal drawings with approval by the representative of buyer (customer) or owner shall be considered the correct measurement and method for fabrication. Delivery schedule will be based on receipt of final approved submittal drawings.\n\nThank you for the opportunity to submit a proposal.\n\nSincerely,")
    document.add_picture('JAG_signature.png', width=Inches(2))

    sign = document.add_paragraph()
    sign.style = document.styles['Normal']
    sign.add_run('Jeff Garlitz\n')
    sign.add_run('jgarlitz@precisionrail.com\n')
    sign.add_run('541-279-8182\n')
    sign.add_run('This estimate was computer generated by Brandon Brodrick').italic = True
    sign.add_run('\n\n\n')
    sign.add_run('Acceptance of Proposal Signature _______________________              Date_______________   ')

    document.save('{} - rev 0.docx'.format(job_name))
    print('proposal finished!')
        
        
    
